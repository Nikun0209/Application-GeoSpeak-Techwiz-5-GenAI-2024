import streamlit as st
import io
import os
import chardet
import fitz  # PyMuPDF
import requests
import hashlib
import fitz # type: ignore
import google.generativeai as gen_ai # type: ignore

from datetime import datetime
from langdetect import detect, DetectorFactory, LangDetectException
from google_gemini import google_gemini_translate, api_key, api_url, translate_role_for_streamlit
from deep_translator import GoogleTranslator
from docx import Document # type: ignore
from dotenv import load_dotenv # type: ignore
from streamlit_option_menu import option_menu
from stability import stability_api_key   
import stability_sdk.interfaces.gooseai.generation.generation_pb2 as generation
from stability_sdk import client
from PIL import Image
from pdf2docx import Converter
from tempfile import NamedTemporaryFile

# Ensure consistency in language detection
DetectorFactory.seed = 0

st.set_page_config(
    page_title="The Avengers - AI Translator",
    page_icon="🤖",
    layout="centered"
)

st.title("🤖 The Avengers - AI Translator")

# Use CSS to position the small subtitle text
st.markdown("""
    <style>
    .subtitle {
        position: absolute;
        top: -.65rem; 
        left: 4.55rem; 
        font-size: 1rem;
        color: gray;
    }
    </style>
            
    <div class="subtitle">Capture semantic meaning!</div>
    <hr>
            
    """, unsafe_allow_html=True)

# Create sidebar
with st.sidebar:
    selected =  option_menu(
        menu_title="Menu",
        options=["Documents", "Text", "ChatBot", "Stability", "PDF to Word", "PDF to PNG", "Blog", "About Us"],
        icons=["file-earmark-text", "alphabet", "robot", "bounding-box", "file-word-fill", "file-pdf-fill", "book", "lightbulb-fill"],
        menu_icon="menu-up",
        default_index=0,
        # orientation="horizontal"
    )

# Tab 1: Document translation
if selected == "Documents":
    col1, col2 = st.columns(2)

    # Select output language
    with col1:
        output_languages_list = [
            'Chinese Simplified', 'Chinese Traditional', 'English', 
            'Vietnamese', 'Korean', 'Japanese'
        ]

        output_language_tab1 = st.selectbox(label="Output language", options=output_languages_list, key="output_language_tab1")

    # File uploader
    with col2:
        uploaded_file = st.file_uploader("Supported file types: .txt, .docx, .pdf", type=["txt", "docx", "pdf"])

    # Add HTML and CSS to change the content
    st.markdown(
        """
        <style>
            .e1b2p2ww11 .e1bju1570::before {
                content: "Please note: We only accept text up to 1000 characters long: .txt, .docx, .pdf";
                font-size: 12px;
                color: white;
                font-family: 'Source Sans Pro', sans-serif;
                letter-spacing: 1px;
            }

            .e1b2p2ww11 .e1bju1570 {
                font-size: 0;
            }
        </style>
        """, unsafe_allow_html=True
    )

    if uploaded_file:
        if st.button("Translate",  key="translate_button_tab1"):
            if uploaded_file.name.endswith('.txt'):
                # Read file as binary
                file_bytes = uploaded_file.read()
                
                # Detect file encoding
                result = chardet.detect(file_bytes)
                encoding = result['encoding']
                
                # Decode file content using detected encoding
                file_contents = file_bytes.decode(encoding)

                # Limit the number of characters
                if len(file_contents) > 1000:
                    st.warning("Limit is 1000 characters.")
                    st.stop()  # Stop the program
                else:        
                    source_language = detect(file_contents)
                    language_mapping = {
                        "Chinese Simplified": "zh-CN",
                        "Chinese Traditional": "zh-TW",
                        "English": "en",
                        "Vietnamese": "vi",
                        "Korean": "ko",
                        "Japanese": "ja"
                    }

                    target_language_code = language_mapping.get(output_language_tab1)

                    if source_language == target_language_code:
                        st.warning(f"The content you provided already in {output_language_tab1}.")
                        st.stop()  # Stop the program
                    else:
                        translated_text = google_gemini_translate(file_contents, None, target_language_code)
                        st.session_state.translated_text_tab1 = translated_text
            elif uploaded_file.name.endswith('.docx'):
                doc = Document(uploaded_file)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                text = '\n'.join(full_text)

                # Limit the number of characters
                if len(text) > 1000:
                    st.warning("Limit is 1000 characters.")
                    st.stop()  # Stop the program
                else:
                    source_language = detect(text)
                    language_mapping = {
                        "Chinese Simplified": "zh-CN",
                        "Chinese Traditional": "zh-TW",
                        "English": "en",
                        "Vietnamese": "vi",
                        "Korean": "ko",
                        "Japanese": "ja"
                    }
                    target_language_code = language_mapping.get(output_language_tab1)
                    
                    if source_language == target_language_code:
                        st.warning(f"The content you provided is already in {output_language_tab1}.")
                        st.stop()  # Stop the program
                    else:
                        input_language = language_mapping.get(source_language)
                        output_language = language_mapping.get(target_language_code)

                        # Translate text using Google Translate API
                        translator = GoogleTranslator(source=source_language, target=target_language_code)
                        translated_text = translator.translate(text)
                        st.session_state.translated_text_tab1 = translated_text
                        # st.session_state.uploaded_file_type = 'docx'
            
            elif uploaded_file.name.endswith('.pdf'):
                # Extract text from the PDF
                with open("temp.pdf", "wb") as f:
                    f.write(uploaded_file.getbuffer())

                doc = fitz.open("temp.pdf")
                full_text = ""
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    full_text += page.get_text()

                # Limit the number of characters
                if len(full_text) > 1000:
                    st.warning("Limit is 1000 characters.")
                    st.stop()  # Stop the program
                else:
                    source_language = detect(full_text)
                    language_mapping = {
                        "Chinese Simplified": "zh-CN",
                        "Chinese Traditional": "zh-TW",
                        "English": "en",
                        "Vietnamese": "vi",
                        "Korean": "ko",
                        "Japanese": "ja"
                    }
                    target_language_code = language_mapping.get(output_language_tab1)

                    if source_language == target_language_code:
                        st.warning(f"The content you provided is already in {output_language_tab1}.")
                        st.stop()  # Stop the program
                    else:
                        input_language = language_mapping.get(source_language)
                        output_language = language_mapping.get(target_language_code)

                        translator = GoogleTranslator(source=source_language, target=target_language_code)
                        translated_text = translator.translate(full_text)
                        st.session_state.translated_text_tab1 = translated_text
                
            else:   
                st.error("Please upload a file with the extension .txt, .docx, .pdf")

    # Display translated result
    if "translated_text_tab1" in st.session_state:
        if uploaded_file.name.endswith('.txt'):
            st.write("Translated content")
            st.success(st.session_state.translated_text_tab1)

            # Create file name based on current time hashed with MD5
            current_time = datetime.now().isoformat()
            md5_hash = hashlib.md5(current_time.encode()).hexdigest()
            file_name = f"{md5_hash}.txt"

            # Create .txt file for translated content and provide download link
            translated_text = st.session_state.translated_text_tab1
            buffer = io.BytesIO()
            buffer.write(translated_text.encode('utf-8'))
            buffer.seek(0)
            
            st.download_button(
                label="Download translation",
                data=buffer,
                file_name=file_name,
                mime="text/plain"
            )
        elif uploaded_file.name.endswith('.docx'):
            st.write("Translated content")
            st.success(st.session_state.translated_text_tab1)
            buffer = io.BytesIO()
            current_time = datetime.now().isoformat()
            md5_hash = hashlib.md5(current_time.encode()).hexdigest()
            file_name = f"{md5_hash}.docx"
            translated_text = st.session_state.translated_text_tab1
            doc = Document()
            doc.add_paragraph(translated_text)
            doc.save(buffer)
            mime_type = "application/msword"

            buffer.seek(0)
            st.download_button(
                label="Download translation",
                data=buffer,
                file_name=file_name,
                mime=mime_type
            )
        elif uploaded_file.name.endswith('.pdf'):
            st.write("Translated content")
            st.success(st.session_state.translated_text_tab1)

            buffer = io.BytesIO()
            current_time = datetime.now().isoformat()
            md5_hash = hashlib.md5(current_time.encode()).hexdigest()
            file_name = f"{md5_hash}.docx"
            translated_text = st.session_state.translated_text_tab1
            doc = Document()
            doc.add_paragraph(translated_text)
            doc.save(buffer)
            mime_type = "application/msword"

            buffer.seek(0)
            st.download_button(
                label="Download translation",
                data=buffer,
                file_name=file_name,
                mime=mime_type
            )
        else:
            st.write("Error")

# Tab 2: Text translation
if selected == "Text":
    col_1, col_2 = st.columns(2)

    # Column for selecting input language
    with col_1:
        input_languages_list = [
            'Chinese Simplified', 'Chinese Traditional', 'English', 
            'Vietnamese', 'Korean', 'Japanese'
        ]
        input_language_tab2 = st.selectbox(label="Input language", options=input_languages_list, key="input_language_tab2")

    # Column for selecting output language
    with col_2:
        output_languages_list = [x for x in input_languages_list if x != input_language_tab2]
        output_language_tab2 = st.selectbox(label="Output language", options=output_languages_list)

    input_text_tab2 = st.text_area("Enter text here (up to 1000 characters)", key="input_text_tab2")

    if st.button("Translate"):
        if input_text_tab2.strip():
            if len(input_text_tab2) > 1000:
                st.warning("Limit is 1000 characters.")
                st.stop()  # Stop the program
            else:
                language_mapping = {
                    "Chinese Simplified": "zh-CN",
                    "Chinese Traditional": "zh-TW",
                    "English": "en",
                    "Vietnamese": "vi",
                    "Korean": "ko",
                    "Japanese": "ja"
                }              
                input_language = language_mapping.get(input_language_tab2)
                output_language = language_mapping.get(output_language_tab2)

                # Translate text using Google Translate API
                translator = GoogleTranslator(source=input_language, target=output_language)
                translated_text = translator.translate(input_text_tab2)
                st.session_state.translated_text_tab2 = translated_text
        else:
            st.warning("Please enter the text to be translated.")

    # Display translated result
    if "translated_text_tab2" in st.session_state:
        st.success(st.session_state.translated_text_tab2)

# Tab 3: ChatBot    
if selected == "ChatBot":
    # Load environment variables
    load_dotenv()

    # Set up Google Gemini-Pro AI model
    gen_ai.configure(api_key=api_key)
    model = gen_ai.GenerativeModel('gemini-pro')

    # Initialize chat session in Streamlit if not already present
    if "chat_session" not in st.session_state:
        st.session_state.chat_session = model.start_chat(history=[])

    # Display the chat history
    for message in st.session_state.chat_session.history:
        with st.chat_message(translate_role_for_streamlit(message.role)):
            st.markdown(message.parts[0].text)

    user_prompt = st.chat_input("Enter your question here...")

    if user_prompt:
        # Add user's message to chat and display it
        st.chat_message("user").markdown(user_prompt)

        try:
            # Send user's message to Gemini-Pro and get the response
            gemini_response = st.session_state.chat_session.send_message(user_prompt)

            # Display Gemini-Pro's response
            with st.chat_message("assistant"):
                st.markdown(gemini_response.text)

        except Exception as e:
            # Catch any exception and display an error message
            st.error(f"An error occurred: {str(e)}")

# Tab 4: Stability
if selected == "Stability":
    # Get the image description from the user
    prompt = st.chat_input("Describe an image you want to create")
    
    if prompt:  # Check if the user has entered a description
        try:
            # Detect the language of the prompt
            language = detect(prompt)
            
            if language == "en":
                # Initialize the Stability AI client
                stability_api = client.StabilityInference(
                    key=stability_api_key,  # Your API key
                    verbose=True,
                )

                # Send the request to generate an image
                answers = stability_api.generate(
                    prompt=prompt,
                    seed=12345,  # Optional: Set a seed to create reproducible results
                    steps=50,  # Number of steps for image creation
                )

                # Process the response from Stability AI
                for resp in answers:
                    for artifact in resp.artifacts:
                        if artifact.finish_reason == generation.FILTER:
                            st.error("Request filtered due to NSFW content")
                        if artifact.type == generation.ARTIFACT_IMAGE:
                            # Open the image from binary data and save it to disk
                            img = Image.open(io.BytesIO(artifact.binary))

                            # Display the image on the web interface
                            st.image(img, caption="The image is created from your description")
            else:
                st.error("Please enter the description in English!")
        except LangDetectException:
            st.error("Could not detect the language. Please enter a valid text!")
    else:
        st.error("Please enter an image description!")

# Tab 5: PDF to Word
if selected == "PDF to Word":
    # Upload PDF file from user
    uploaded_file = st.file_uploader("Convert PDF to Word", type=["pdf"])

    if uploaded_file is not None:
        original_filename = os.path.splitext(uploaded_file.name)[0]

        # Create a temporary file to save the uploaded file
        with NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(uploaded_file.getbuffer())
            temp_pdf_path = temp_pdf.name

        # Convert from PDF to Word
        cv = Converter(temp_pdf_path)
        output_file = "output.docx"
        cv.convert(output_file, start=0, end=None)
        cv.close()

        # Display a success message and provide a download link for the Word file
        st.success("Conversion successful!")
        
        new_filename = f"{original_filename}_the_avengers.docx"

        # Create a button to download the Word file
        with open(output_file, "rb") as f:
            download_button = st.download_button(
                label="Download",
                data=f,
                file_name=new_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # Check if the user has clicked the download button
        if download_button:
            try:
                if os.path.exists(temp_pdf_path):
                    os.remove(temp_pdf_path)
                if os.path.exists(output_file):
                    os.remove(output_file)
            except Exception as e:
                st.error(f"An error occurred while deleting the files: {e}")

# Tab 6: PDF to PNG
if selected == "PDF to PNG":
    # Upload PDF file from the user
    uploaded_file = st.file_uploader("Convert PDF to PNG", type=["pdf"])

    if uploaded_file is not None:
        # Get the original file name without the extension
        original_file_name = uploaded_file.name.rsplit('.', 1)[0]

        # Create a temporary file to save the uploaded file
        with NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(uploaded_file.getbuffer())
            temp_pdf_path = temp_pdf.name

        # Open the PDF file and convert each page to PNG
        pdf_document = fitz.open(temp_pdf_path)
        png_files = []

        for page_number in range(len(pdf_document)):
            page = pdf_document.load_page(page_number)
            pix = page.get_pixmap()

            # Convert the pixmap to a PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Determine the output PNG file name
            if len(pdf_document) == 1:
                output_png_name = f"{original_file_name}_the_avengers.png"
            else:
                output_png_name = f"{original_file_name}_the_avengers_page_{page_number + 1}.png"

            # Save the image to an in-memory byte stream
            png_io = io.BytesIO()
            img.save(png_io, format="PNG")
            png_io.seek(0)  # Rewind to the start of the byte stream
            png_files.append((output_png_name, png_io))

        pdf_document.close()

        # Display success message and provide download links for PNG images
        st.success("Conversion successful!")

        # Create download buttons for PNG images from memory
        for png_name, png_io in png_files:
            st.download_button(
                label=f"Download {png_name}",
                data=png_io,
                file_name=png_name,
                mime="image/png"
            )

        # Delete the temporary PDF file
        try:
            if os.path.exists(temp_pdf_path):
                os.remove(temp_pdf_path)
        except Exception as e:
            st.error(f"An error occurred while deleting files: {e}")

# Tab 7: Blog
if selected == "Blog":
     # st.title("Techwiz 5 - GeoSpeak - Developed by The Avengers")
    st.title("Techwiz 5 - 2024 - Global IT Competition - 43 nations - over 810 teams")
    # st.subheader("Leveraging Gen AI for Smart solution of Translation: Geo-Speak Application")

    st.markdown(
    """
        Category: GenAI \n
        Project: GeoSpeak - AI Translator & Chatbot \n

        Lecturer: MSc. Hồ Nhựt Minh
        
        **THE AVENGERS - AI TRANSLATOR PROJECT MEMBERS**
        - Dương Gia Thành: Main Developer
        - Nguyễn Quốc Anh: Project Manager
        - Nguyễn Quốc Bảo
        - Phạm Hoàng Tiến
        - Trần Mai Phương
        - Lâm Kim Khánh

        Website: [Visit here](https://theavengers.streamlit.app/)\n
        User guide: [Visit here](https://youtu.be/6ZnQhEFjJZ8)
    """)

    st.image(
        "img/z5855806767554_aef51eacda4c36fd65660ce9ee04af64.jpg", 
        width=50, 
        caption="The Avengers Team", 
        use_container_width=True
    )

    st.markdown(
    """
        **What is going on with Translation?**

        <p style="text-align: justify;">In the fast pace of world, it is common that people are breathing the cosmopolitan atmosphere, even companies. What does it mean? People have engaged in using in diverse languages more than ever before in their workplaces. They utilize languages or linguistic skills in reading and writing documents, listening to foreigners and especially speaking in a manner of responding actively to any information that is not of their “mother tongue” they received. To this point, some may think of many conventional translation tools as a solution which automatically translate any text that they input. It is fair enough to agree, but, the main point is that those tools are just simply doing the job of translation.</p>

        <p style="text-align: justify;">Take a closer look at the problem mentioned there, the “job” those conventional translation tools do is transferring words people input into a targeted language form without any consideration of context. It is just like users search for a vocabulary in a thick dictionary book and then use it, but maybe they do not understand how the word can be effectively use in the context they mean. It is about the problem of word choices. Those conventional tools may translate text accurately; however, it would cause difficulties of understanding the text or even misunderstandings. When it turns into the business context in which accuracy is the priority, the problem starts to arise.  People may find it hard to read many technical documents or even communicate about technical problems with many foreign experts. The sufficient flow of information exchange can be a result afterwards. One of problems related to reading technical documents of computer science fields, it is obvious that the same word can have different meanings.</p>

        <p style="text-align: justify;">Given those problems, we, human-beings, starts to think of something smarter and more sophisticated in providing translation service with a good combination of contextual awareness and linguistic profession. Therefore, Gemini's Large language Model and Generative AI (Gen AI) are utilized in developing a translation solution, called Geo-Speak.</p>

        **What is Geo-Speak?**

        <p style="text-align: justify;">To tackle issues in the language use, an application, called Geo-Speak, is innovated, aiming to address the limitations of traditional translation methods and the increasing complexity of language usage across diverse domains. By applying Gemini's Large Language Models (LLMs) and Generative AI (Gen AI), Geo-Speak are likely to interpret and translate text in manners of combining the contextual factors and academic linguistics. This application is essential for diplomacy, global business, education, and so on, where precise communication can significantly affect the outcomes. Hence, users are able to overcome language barriers, facilitate cross-cultural understanding, and ensure the smooth flow of information exchanging regardless of personal background or linguistic profession.</p>
        
    """, unsafe_allow_html=True)
    
    st.image("img/z5855785979880_853b682305e933c24c9e8eae5c40f52e.jpg", width=50, caption="Architecture", use_container_width=True)

    st.markdown(
    """
        <p style="text-align: justify;">Geo-Speak is a sophisticated real-time web application providing translating services. It is worth to mention that its design is the utilization of cutting-edge AI technologies, particularly a cloud-based Large Language Model, to deliver effective and efficient communication across diverse languages. Not only it enables users overcome the language barriers, but also take contextual sensitiveness into consideration to ensure the flow of information exchanging. It is fair to say that Geo-Speak is a human-kind leap for a context in which people have to adapt to future expansion involved in multinational and multilingual working environments quickly. By leveraging Gemini’s models, which specializes in language comprehension and generation. This solution offers real-time translations that are more accurate and contextually sensitive compared to the conventional methods. Therefore, Geo-Speak can potentially enhance cross-language communication in diverse workplace settings, being indispensable in businesses. </p>

        **How the Geo-Speak Application works?**

        **There are two main phrases performed by the Geo-Speak application:**

        <p style="text-align: justify;">At first, users inputs text or text file (.txt, .docx, .pdf extension) in any language forms and choose the desired target language from a drop-down list provided in the web interface. The Embedding API accesses the input information. At this stage, the application use constraints, related to word limitations and extension of a file to check if the file or text is qualified enough to process the further steps. Particularly, in this project, my team set 1000 words as the limitation and “.txt, .docx, .pdf” extension for the standard of file format. Once verified, the input information will be decoded into a high-dimensional embedding vector that captures its semantic essence. This vector is then used to query a vector Database containing precomputed embeddings of parallel corporate or translation sample. The system extracts relevant documents that closely aligned with the source text embedding, comprising paired texts in both source and target languages to guide the translation process. </p>
    """, unsafe_allow_html=True)

    st.image("img/z5855786203327_f6ec2b25a2b6acbed2fdbaabd8c6de88.jpg", width=50, caption="User Journey Description", use_container_width=True)

    st.markdown(
    """
        <p style="text-align: justify;">In the second phrase, the application transfers those data to Google Gemini API for translating process. The translated text is sent back to the users afterwards in a text or text file which is available for reading or downloading.</p>

        **Technologies and Tools in building Geo-Speak**

        <p style="text-align: justify;">To build Geo-Speak application, we start from scratches that maybe most people find themselves common with terms.</p>
    """, unsafe_allow_html=True)

    st.image("img/z5855785794142_784584c673ba215f438c3d5329e4ed59.jpg", width=50, caption="Used Technologies", use_container_width=True)

    st.markdown(
    """
        <p style="text-align: justify;">For the use of Programming Languages, Python would be smart options. Python will do the jobs of backend development, data processing and integration with AI and natural language toolkit libraries. Meanwhile, HTML and CSS is used for jobs of front-end development, interactivity and user interface design. Relatively, Frameworks of Python can be used along with developing responsive and interactive user interfaces.</p>

        <p style="text-align: justify;">For User Interface Design, developers can use HTML and CSS for building responsive and visually appealing user interface of website.</p>

        <p style="text-align: justify;">For Translation APIs, we can utilize which already have before, such as Google Cloud Translation API, Google Gemini API, for integrating neural machine translation capabilities into the application.</p>

        **Real-world Applications of Geo-Speak in business communications**

        <p style="text-align: justify;">Above is all you need to know about Geo-Speak application, It is time to look at this tool from business aspects to have a comprehensive understanding of the reasons we may need it.</p>
    """, unsafe_allow_html=True)

    st.image("img/tumblr_d5dff37917648d2f077d1971bc3335ce_8c896f3e_500.jpg", width=50, caption="Application GeoSpeak", use_container_width=True)

    st.markdown(
    """
        <p style="text-align: justify;">In E-commerce industry, Online retailers definitely want to expand their brand name and products across countries over the world. The presence of Geo-Speak is useful for providing multilingual customers with sophisticated supports, product descriptions, and also checkout processes. From there, it fosters the better shopping experience.</p>

        <p style="text-align: justify;">For Travel and tourism industry, it is not necessary to explain the role of language anymore. Travel agencies, local businesses, like accommodation, restaurants, or tour guide services need to be specialized in using languages naturally more than any industries. Geo-Speak would help in offering real-time language assistance throughout the customer journeys.</p>

        <p style="text-align: justify;">For Legal and Compliance areas, corporation have to deal with laws across borders. Speaking of this context, strict word choices and comprehensive understandings of words’ meanings would form language barriers for them when they want to expand oversea. Therefore, Geo-Speak with high level of translation accuracy can facilitate the process of conducting documents, contracts, and compliance materials effectively.</p>

        <p style="text-align: justify;">There are still many areas to be mentioned to demonstrate the usefulness of this sophisticated Application.</p>
    """, unsafe_allow_html=True)

# Tab 8: About Us
if selected == "About Us":
    col1, col2, col3 = st.columns(3)

    with col2:
        st.image(
            "img/z5855663862790_aa6a77c21ae95ad7f83fb3eccb262482.jpg", 
            width=50,  
            caption="The Avengers Team", 
            use_container_width=True
        )

    st.write(
        """
        **ABOUT US** \n
        Techwiz 5 - 2024 - Global IT Competition - 43 nations - over 810 teams \n
        Category: GenAI \n
        Project: GeoSpeak - AI Translator & Chatbot \n

        Lecturer: MSc. Hồ Nhựt Minh
        
        **THE AVENGERS - AI TRANSLATOR PROJECT MEMBERS**
        - Dương Gia Thành: Main Developer
        - Nguyễn Quốc Anh: Project Manager
        - Nguyễn Quốc Bảo
        - Phạm Hoàng Tiến
        - Trần Mai Phương
        - Lâm Kim Khánh

        Website: [Visit here](https://theavengers.streamlit.app/)\n

        **CONTACT**
        - Dương Gia Thành: [duonggiathanh3819@gmail.com](mailto:duonggiathanh3819@gmail.com)
        - Nguyễn Quốc Anh: [anh.datascience@gmail.com](mailto:anh.datascience@gmail.com)
        """
    )
